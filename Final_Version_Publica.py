from docx import Document
import pandas as pd
from docx2pdf import convert
import os
import datetime

# Obtén la fecha actual
fecha_actual = datetime.datetime.now()

# Formatea la fecha como 'dia_mes_año'
nombre_carpeta = fecha_actual.strftime("%d_%m_%Y")

# Obtén la ruta del directorio actual
ruta_actual = os.getcwd()

# Combina la ruta actual con el nombre de la carpeta
ruta_completa = os.path.join(ruta_actual, nombre_carpeta)

def check_value(val, positivo=[False, 'SI', 'PERSONA_MORAL']):
    if val in positivo:
        return "X", ""
    else:
        return "", "X"
nombre_t = "declaraciones_publicas.csv"
ruta_t = os.path.join(ruta_completa, nombre_t)
df=pd.read_csv(ruta_t)
#df = pd.read_csv("C:\\Users\\vladi\\OneDrive\\Desktop\\Declaraciones publicas\\declaraciones_publicas.csv")
df = df[df['encabezado_tipoDeclaracion'] != 'NOTA']
df = df.reset_index(drop=True)


for index, row in df.iterrows():
    nombre_doc = "VERSION PUBLICA DE  DECLARACION PATRIMONIAL.docx"
    ruta_do= os.path.join(ruta_actual, nombre_doc)
    doc = Document(ruta_do)
    
    bancoSi, bancoNo = check_value(row["declaracion_inversionesCuentasValores_ninguno"])
    bienesSi, bienesNo = check_value(row["declaracion_bienesInmuebles_ninguno"])
    vehiculosSi, vehiculosNo = check_value(row["declaracion_vehiculos_ninguno"])
    extranjeroSi, extranjeroNo = check_value(row["domicilioExtranjero"])
    moralSi, moralNo = check_value(row["moral_o_fisica"])
    ingresHogarSi, ingresHogarNo = check_value(row["declaracion_ingresos_ingresoNetoParejaDependiente_remuneracion_monto"])
    otroSi, otroNo = check_value(row["declaracion_ingresos_otrosIngresosTotal_monto"])
    dependientSi, dependientNo = check_value(row["declaracion_datosDependientesEconomicos_ninguno"])
    # Crear un diccionario para los reemplazos en párrafos
    reemplazos_parrafos = {
        "{Tipo}": row["encabezado_tipoDeclaracion"],
        "{Fecha}": row["encabezado_fechaActualizacion"],
        "{Folio}": row["_id"]
    }

    # Crear un diccionario para los reemplazos en tablas
    print(row["nombre"])
    print(row["areaAdscripcion"], row["empleoCargoComision"], row["fechaEncargo"])
    reemplazos_tablas = {
        "{Nombre}": str(row["nombre"]),
        "{Area}": str(row["areaAdscripcion"]),
        "{Empleo}": str(row["empleoCargoComision"]),
        "{FechaIngreso}": str(row["fechaEncargo"]),
        "{Grado}": str(row["max_nivel_academico"]),
        "{BancoSi}": bancoSi, 
        "{BancoNo}": bancoNo,
        "{BienesSi}": bienesSi, 
        "{BienesNo}": bienesNo,
        "{VehiculosSi}": vehiculosSi,
        "{VehiculosNo}": vehiculosNo,
        "{ExtranjeroSi}": extranjeroSi, 
        "{ExtranjeroNo}": extranjeroNo,
        "{MoralSi}": moralSi, 
        "{MoralNo}": moralNo,
        "{IngresHogarSi}": ingresHogarSi, 
        "{IngresHogarNo}": ingresHogarNo,
        "{OtroSi}": otroSi, 
        "{OtroNo}": otroNo,
        "{DependientSi}": dependientSi,
        "{DependientNo}": dependientNo
    }

    # Reemplazo en párrafos
    for paragraph in doc.paragraphs:
        for marcador, reemplazo in reemplazos_parrafos.items():
            if marcador in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(marcador, reemplazo)

    # Reemplazo en tablas
    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                for marcador, reemplazo in reemplazos_tablas.items():
                    if marcador in cell.text:
                        cell.text = cell.text.replace(marcador, reemplazo)



    # Guardar el documento .docx en la carpeta creada
    ruta_docx = os.path.join(ruta_completa, f"{row['nombre']}_{row['encabezado_tipoDeclaracion']}.docx")
    print(ruta_docx)
    doc.save(ruta_docx)

    # Convertir a PDF y guardar en la misma carpeta
    ruta_pdf = os.path.join(ruta_completa, f"{row['nombre']}_{row['encabezado_tipoDeclaracion']}.pdf")
    convert(ruta_docx, ruta_pdf)

    # Guardar el documento con el nombre de la fila en la columna "nombre"

    #doc.save(f"Declaraciones_despues_29_08\\{row['nombre']}_{row['encabezado_tipoDeclaracion']}.docx")

    #convert(f"C:\\Users\\vladi\\OneDrive\\Desktop\\Declaraciones publicas\\{row['nombre']}.docx")
    #convert(f"Declaraciones_despues_29_08\\{row['nombre']}_{row['encabezado_tipoDeclaracion']}.docx", f"Declaraciones\\{row['nombre']}_{row['encabezado_tipoDeclaracion']}.pdf")